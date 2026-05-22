"""Build a comprehensive technical reference Word document.

Run from project root:
    source venv/bin/activate
    python scripts/build_technical_reference.py

Output: docs/AI_Expense_Tracker_Technical_Reference.docx
"""
import ast
import os
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

PROJECT_ROOT = Path(__file__).parent.parent
OUTPUT = PROJECT_ROOT / "docs" / "AI_Expense_Tracker_Technical_Reference.docx"

# ────────────────────────────────────────────────────────────────────────────
# Files we describe in the dependency map (skip venv, .git, docs/, scripts/)
# ────────────────────────────────────────────────────────────────────────────
SOURCE_FILES = [
    "main.py",
    "database.py",
    "models.py",
    "routes/__init__.py",
    "routes/expenses.py",
    "routes/salary.py",
    "routes/ai.py",
    "routes/settings.py",
    "ai/__init__.py",
    "ai/prompts.py",
    "ai/tools.py",
    "ai/llm_client.py",
    "services/__init__.py",
    "services/currency.py",
    "frontend/app.py",
    "frontend/api_client.py",
    "frontend/formatting.py",
    "frontend/pages/1_Monthly_Summary.py",
    "frontend/pages/2_Yearly_Summary.py",
    "frontend/pages/3_AI_Insights.py",
    "frontend/pages/4_Chat.py",
    "frontend/pages/5_Settings.py",
    "requirements.txt",
    "README.md",
    ".gitignore",
]

FILE_ROLES = {
    "main.py": "FastAPI app entry — registers routers + CORS",
    "database.py": "MongoDB connection (Motor async client) + collection handles",
    "models.py": "Pydantic request/response models for the API",
    "routes/__init__.py": "Marks routes/ as a Python package",
    "routes/expenses.py": "Expense CRUD + monthly/yearly summary endpoints",
    "routes/salary.py": "Salary set/add endpoints + savings calculator (with currency conv.)",
    "routes/ai.py": "AI-powered endpoints (smart-add, natural-add, insights, advice, chat) — natural-add prefers AI-detected currency over dropdown",
    "routes/settings.py": "Base-currency setting + bulk conversion endpoint",
    "ai/__init__.py": "Marks ai/ as a Python package",
    "ai/prompts.py": "All LLM system prompts (insights, advice, categorize, parse, chat) — chat prompt injects today's date dynamically",
    "ai/tools.py": "Sync MongoDB tool functions Groq can call (11 tools) — every tool result includes pre-formatted *_display fields",
    "ai/llm_client.py": "Groq client wrapper — single swap point for LLM provider; injects today's date into the chat system prompt",
    "services/__init__.py": "Marks services/ as a Python package",
    "services/currency.py": "Exchange-rate fetcher (frankfurter.dev) + conversion + locale-aware format_amount helper",
    "frontend/app.py": "Streamlit Dashboard — salary, add forms, expense list",
    "frontend/api_client.py": "Thin HTTP wrapper around the FastAPI backend",
    "frontend/formatting.py": "Locale-aware number formatting helper (Indian vs Western)",
    "frontend/pages/1_Monthly_Summary.py": "Streamlit page — month filter + pie chart",
    "frontend/pages/2_Yearly_Summary.py": "Streamlit page — year-to-date bar charts",
    "frontend/pages/3_AI_Insights.py": "Streamlit page — AI summary + budget advice",
    "frontend/pages/4_Chat.py": "Streamlit page — conversational chatbot UI",
    "frontend/pages/5_Settings.py": "Streamlit page — base currency selector",
    "requirements.txt": "Pinned Python dependencies for reproducible installs",
    "README.md": "Public-facing project description for GitHub",
    ".gitignore": "Patterns to exclude from git (venv, .env, secrets)",
}


# ────────────────────────────────────────────────────────────────────────────
# Helpers
# ────────────────────────────────────────────────────────────────────────────
MONO_FONT = "Courier New"


def configure_document(doc: Document):
    section = doc.sections[0]
    section.page_height = Inches(11)
    section.page_width = Inches(8.5)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)


def add_h1(doc, text):
    return doc.add_heading(text, level=1)


def add_h2(doc, text):
    return doc.add_heading(text, level=2)


def add_h3(doc, text):
    return doc.add_heading(text, level=3)


def add_para(doc, text, italic=False, bold=False):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.italic = italic
    r.bold = bold
    return p


def add_code(doc, code: str):
    """Add a monospace code block (Courier New, 9pt)."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.2)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(code)
    r.font.name = MONO_FONT
    r.font.size = Pt(9)
    return p


def add_mixed_para(doc, segments):
    """Add a paragraph with mixed plain + mono runs.
    segments is a list of (text, kind) where kind is 'text' or 'mono'.
    """
    p = doc.add_paragraph()
    for text, kind in segments:
        r = p.add_run(text)
        if kind == "mono":
            r.font.name = MONO_FONT
            r.font.size = Pt(10)
        elif kind == "bold":
            r.bold = True
    return p


def add_bullet(doc, text):
    return doc.add_paragraph(text, style="List Bullet")


def add_table(doc, rows: list[list[str]], header=True, mono_cols=None):
    """Build a table from rows. First row is header if header=True.
    mono_cols: list of column indices to render in Courier New.
    """
    mono_cols = mono_cols or []
    table = doc.add_table(rows=len(rows), cols=len(rows[0]))
    table.style = "Light Grid Accent 1"
    for r_idx, row in enumerate(rows):
        for c_idx, cell_text in enumerate(row):
            cell = table.cell(r_idx, c_idx)
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(cell_text)
            if r_idx == 0 and header:
                run.bold = True
                run.font.size = Pt(10)
            else:
                run.font.size = Pt(9)
            if c_idx in mono_cols and not (r_idx == 0 and header):
                run.font.name = MONO_FONT
    return table


def parse_imports(file_path: Path):
    """Return list of (module, [names]) for a Python file."""
    try:
        src = file_path.read_text()
        tree = ast.parse(src)
    except Exception:
        return []
    imports = []
    for node in ast.walk(tree):
        if isinstance(node, ast.Import):
            for alias in node.names:
                imports.append((alias.name, []))
        elif isinstance(node, ast.ImportFrom):
            module = node.module or ""
            names = [a.name for a in node.names]
            imports.append((module, names))
    return imports


def project_relative_imports(file_relpath: str) -> list[str]:
    """Return the list of OTHER project files this file imports."""
    file_path = PROJECT_ROOT / file_relpath
    if not file_path.suffix == ".py":
        return []
    raw_imports = parse_imports(file_path)
    out = []
    for module, names in raw_imports:
        if not module:
            for n in names:
                if any(n.startswith(p) for p in ("routes", "ai", "services", "frontend")):
                    out.append(n)
            continue
        # Map module name to project file
        candidates = [
            module.replace(".", "/") + ".py",
            module.replace(".", "/") + "/__init__.py",
        ]
        for c in candidates:
            if (PROJECT_ROOT / c).exists():
                out.append(c)
                break
    return sorted(set(out))


def build_dependency_data():
    """Return dict: file -> {imports: [], imported_by: []}."""
    data = {f: {"imports": [], "imported_by": []} for f in SOURCE_FILES}
    for f in SOURCE_FILES:
        imps = project_relative_imports(f)
        data[f]["imports"] = imps
        for imp in imps:
            if imp in data:
                data[imp]["imported_by"].append(f)
    return data


# ────────────────────────────────────────────────────────────────────────────
# Section builders
# ────────────────────────────────────────────────────────────────────────────


def cover_page(doc):
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for _ in range(8):
        doc.add_paragraph()

    h = doc.add_paragraph()
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = h.add_run("AI Expense Tracker")
    run.font.size = Pt(36)
    run.bold = True

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sr = sub.add_run("Personal Technical Reference")
    sr.font.size = Pt(18)
    sr.italic = True

    for _ in range(8):
        doc.add_paragraph()

    d = doc.add_paragraph()
    d.alignment = WD_ALIGN_PARAGRAPH.CENTER
    dr = d.add_run(f"Generated: {date.today().strftime('%B %d, %Y')}")
    dr.font.size = Pt(12)

    author = doc.add_paragraph()
    author.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ar = author.add_run("Author: Hiya Arora")
    ar.font.size = Pt(12)

    doc.add_page_break()


def section_1_project_overview(doc):
    add_h1(doc, "1. Project Overview")

    add_para(doc,
        "The AI Expense Tracker is a full-stack personal-finance application built "
        "as a portfolio project. It allows the user to log expenses, set salary, view "
        "monthly/yearly spending breakdowns and chat with an AI assistant that can both "
        "read and modify the underlying database through natural-language conversation. "
        "Multi-currency support is built in: any expense or salary entered in a foreign "
        "currency is auto-converted to the user's base currency using live exchange rates."
    )

    add_para(doc,
        "The project intentionally demonstrates three modern LLM-engineering patterns: "
        "LLM-as-Classifier (smart category picker), Structured Data Extraction "
        "(parsing free-form sentences into JSON) and Tool Use / Function Calling "
        "(the chatbot)."
    )

    add_h2(doc, "Tech Stack")

    tech_rows = [
        ["Layer", "Tool", "Purpose"],
        ["Backend framework", "FastAPI", "Turns Python functions into REST endpoints"],
        ["Async DB client", "Motor", "Async MongoDB client used inside FastAPI"],
        ["Sync DB client", "PyMongo", "Sync MongoDB client used inside AI tools"],
        ["Database", "MongoDB Atlas", "Document-store for expenses, salaries, settings"],
        ["LLM provider", "Groq", "Hosts the openai/gpt-oss-120b model"],
        ["LLM model", "openai/gpt-oss-120b", "Open-source GPT model tuned for tool calling"],
        ["LLM SDK", "groq (Python)", "OpenAI-compatible tool-calling SDK"],
        ["Validation", "Pydantic v2", "Models incoming JSON + auto-docs"],
        ["Frontend", "Streamlit", "Python-only multi-page web UI"],
        ["Charts", "Plotly Express", "Interactive pie + bar charts"],
        ["HTTP client", "httpx", "Calls FastAPI from Streamlit"],
        ["Exchange rates", "frankfurter.dev", "Free, key-less FX API"],
        ["Env vars", "python-dotenv", "Loads .env at startup"],
        ["Hosting target", "Railway + Streamlit Cloud", "Backend + frontend"],
    ]
    add_table(doc, tech_rows, header=True, mono_cols=[1])

    doc.add_page_break()


def section_2_entry_points(doc):
    add_h1(doc, "2. Entry Points")

    # ---- Backend ----
    add_h2(doc, "Backend entry point: main.py")

    add_para(doc,
        "The backend process is started with the command:"
    )
    add_code(doc, "uvicorn main:app --reload")

    add_para(doc,
        "Uvicorn is an ASGI (Asynchronous Server Gateway Interface) server. The argument "
        "'main:app' tells Uvicorn to import the variable 'app' from the file 'main.py'."
    )
    add_para(doc, "Contents of main.py:")
    add_code(doc, (PROJECT_ROOT / "main.py").read_text())

    add_h3(doc, "What happens at startup")
    add_bullet(doc, "FastAPI() creates an empty application object.")
    add_bullet(doc, "CORSMiddleware is wrapped around the app so the Streamlit frontend (a different origin) is allowed to call it.")
    add_bullet(doc, "Each include_router(...) call walks through the routes module, finds every @router.get/post/put/delete handler, and registers it on the app's URL map.")
    add_bullet(doc, "Uvicorn opens TCP port 8000 and starts accepting HTTP requests.")
    add_bullet(doc, "Side effects from imports: database.py runs AsyncIOMotorClient(...) at import time, opening a persistent connection to MongoDB Atlas; ai/llm_client.py instantiates a Groq() client; services/currency.py initialises an in-memory rate cache; ai/tools.py opens a sync PyMongo connection.")

    # ---- Frontend ----
    add_h2(doc, "Frontend entry point: frontend/app.py")
    add_para(doc, "The frontend process is started with:")
    add_code(doc, "streamlit run frontend/app.py")

    add_para(doc,
        "Streamlit treats this file as the 'home page' of a multi-page app. Any *.py file "
        "under frontend/pages/ automatically becomes another page in the sidebar."
    )

    add_h3(doc, "What happens at startup")
    add_bullet(doc, "Streamlit starts a Tornado web server on port 8501.")
    add_bullet(doc, "When the user opens http://localhost:8501 in a browser, Streamlit runs frontend/app.py top-to-bottom from a fresh Python process.")
    add_bullet(doc, "frontend/app.py calls get('/settings') on the FastAPI backend to fetch the user's base currency and supported codes, used to render currency symbols.")
    add_bullet(doc, "Each Streamlit widget (st.text_input, st.form, etc.) renders an HTML/JS element in the browser.")
    add_bullet(doc, "Every user interaction (button click, form submit) triggers a full top-to-bottom rerun of the script — Streamlit's execution model.")

    doc.add_page_break()


def section_3_dependency_map(doc):
    add_h1(doc, "3. File Dependency Map")

    add_para(doc,
        "The table below lists every source file in the project. The 'Imports' column "
        "shows what other project files this file depends on; 'Imported by' shows the reverse. "
        "External-library imports (fastapi, motor, plotly, etc.) are omitted to keep the focus on "
        "the internal architecture."
    )

    deps = build_dependency_data()

    rows = [["File", "Imports (project files)", "Imported by", "Role"]]
    for f in SOURCE_FILES:
        imports_cell = "\n".join(deps[f]["imports"]) or "—"
        imported_by_cell = "\n".join(deps[f]["imported_by"]) or "—"
        role = FILE_ROLES.get(f, "")
        rows.append([f, imports_cell, imported_by_cell, role])

    add_table(doc, rows, header=True, mono_cols=[0, 1, 2])
    doc.add_page_break()


# ────────────────────────────────────────────────────────────────────────────
# Section 4: Data flow — five major flows
# ────────────────────────────────────────────────────────────────────────────


def flow_step(doc, title: str, file_path: str, code: str, data_in: str, data_out: str, narration: str):
    add_h3(doc, title)
    add_mixed_para(doc, [("File: ", "text"), (file_path, "mono")])
    add_para(doc, "Code:")
    add_code(doc, code)
    add_para(doc, "Data shape entering this step:")
    add_code(doc, data_in)
    add_para(doc, "Data shape leaving this step:")
    add_code(doc, data_out)
    add_para(doc, "What happens here:")
    add_para(doc, narration)


def section_4_data_flow(doc):
    add_h1(doc, "4. Data Flow")
    add_para(doc,
        "This section traces five major routes end-to-end. For each step we show the file, "
        "the relevant code, the data shape entering and leaving, and a plain-English "
        "explanation of the transformation."
    )

    # ───────────────────── Flow 1: Natural-Language expense add ──────
    add_h2(doc, "Flow 1: POST /expenses/natural — Natural-language expense add")
    add_para(doc,
        "User scenario: user types 'I spent 1000 on zomato food' in the Dashboard, "
        "Currency = INR, Date = today. The AI parses the sentence into structured fields "
        "and inserts a document into MongoDB."
    )

    flow_step(doc,
        "Step 1.1 — Streamlit form submission",
        "frontend/app.py",
        '''with st.form("nl_form", clear_on_submit=True):
    nl_text = st.text_input("Describe the expense in one line", ...)
    nl_date = st.date_input("Date", value=date.today())
    nl_curr_label = st.selectbox("Currency typed in", CURRENCY_LABELS, ...)
    nl_currency = CURRENCY_CODES[CURRENCY_LABELS.index(nl_curr_label)]
    if st.form_submit_button("Parse & Add with AI"):
        result = post("/expenses/natural", {
            "text": nl_text,
            "date": nl_date.isoformat(),
            "currency": nl_currency,
        })''',
        "(User input, in Streamlit widget memory)\n"
        "  nl_text   : 'I spent 1000 on zomato food'\n"
        "  nl_date   : date(2026, 5, 18)\n"
        "  nl_currency: 'INR'",
        "(HTTP request body)\n"
        "{\n"
        '  "text": "I spent 1000 on zomato food",\n'
        '  "date": "2026-05-18",\n'
        '  "currency": "INR"\n'
        "}",
        "The form button triggers a rerun. Streamlit calls api_client.post which serialises "
        "the dict to JSON and sends it to the FastAPI backend.",
    )

    flow_step(doc,
        "Step 1.2 — api_client.post serialises and sends",
        "frontend/api_client.py",
        '''def post(path: str, payload: dict):
    r = httpx.post(f"{BACKEND_URL}{path}", json=payload, timeout=120.0)
    r.raise_for_status()
    return r.json()''',
        "path    : '/expenses/natural'\n"
        "payload : { ... see above ... }",
        "(HTTP wire-level)\n"
        "POST http://localhost:8000/expenses/natural\n"
        "Content-Type: application/json\n"
        "{...JSON body...}",
        "httpx serialises the Python dict to a JSON byte stream, sets the Content-Type "
        "header and POSTs to the backend.",
    )

    flow_step(doc,
        "Step 1.3 — FastAPI validates with Pydantic",
        "routes/ai.py + models.py",
        '''class NaturalExpense(BaseModel):
    text: str
    date: Optional[str] = None
    currency: Optional[str] = None

@router.post("/expenses/natural")
async def natural_add_expense(nl: NaturalExpense):
    ...''',
        "(Raw JSON body)\n"
        '{"text": "...", "date": "...", "currency": "..."}',
        "(Validated Python object, all types enforced)\n"
        'nl = NaturalExpense(\n'
        '  text="I spent 1000 on zomato food",\n'
        '  date="2026-05-18",\n'
        '  currency="INR"\n'
        ")",
        "Before our handler runs, FastAPI matches the JSON body against the NaturalExpense "
        "model. If a field is missing or has the wrong type, FastAPI returns HTTP 422 "
        "automatically and our function is never called.",
    )

    flow_step(doc,
        "Step 1.4 — AI extracts structured data from text (now incl. currency hint)",
        "ai/llm_client.py",
        '''def parse_expense_text(text: str) -> dict:
    raw = _ask_groq(PARSE_EXPENSE_PROMPT, text, max_tokens=250)
    cleaned = raw.strip()
    if cleaned.startswith("```"):
        cleaned = cleaned.strip("`")
        if cleaned.lower().startswith("json"):
            cleaned = cleaned[4:].strip()
    parsed = json.loads(cleaned)
    raw_curr = parsed.get("currency")
    currency_hint = None
    if isinstance(raw_curr, str) and _supported_currency(raw_curr.upper()):
        currency_hint = raw_curr.upper()
    return {
        "title": parsed["title"].strip(),
        "amount": float(parsed["amount"]),
        "category": parsed["category"] if parsed["category"] in valid_cats else "Miscellaneous",
        "currency_hint": currency_hint,
    }''',
        "text = 'paid 500 rs in shopping mall'",
        "{'title': 'Shopping mall', 'amount': 500.0, 'category': 'Shopping', 'currency_hint': 'INR'}",
        "PARSE_EXPENSE_PROMPT instructs Groq to return EXACTLY 4 keys, including currency. "
        "The model maps natural-language hints (rs, $, €, etc.) to 3-letter codes. "
        "Code-fence stripping and json.loads happen as before. The currency hint is then "
        "validated against the supported set — anything unrecognised becomes None.",
    )

    flow_step(doc,
        "Step 1.5 — Pick the input currency (3-tier priority)",
        "routes/ai.py",
        '''input_currency = (
    parsed.get("currency_hint")
    or (nl.currency or base_currency)
).upper()
try:
    amount_fields = build_amount_fields(parsed["amount"], input_currency, base_currency)
except ValueError as e:
    raise HTTPException(status_code=400, detail=str(e))''',
        "parsed['currency_hint']='INR', nl.currency=None, base_currency='USD'",
        "input_currency = 'INR'  (AI-detected wins)",
        "Priority order: AI-detected currency from the text > explicit dropdown value > "
        "base currency. The build_amount_fields call is now wrapped in try/except so that "
        "an unsupported code (e.g. the literal 'string' placeholder from Swagger UI) returns "
        "a clear HTTP 400 instead of crashing the server with a 500.",
    )

    flow_step(doc,
        "Step 1.6 — Currency conversion to base currency",
        "services/currency.py",
        '''def build_amount_fields(amount, input_currency, base_currency) -> dict:
    input_currency = (input_currency or "").upper()
    if not supported(input_currency):
        raise ValueError(f"Unsupported currency: '{input_currency}'")
    if input_currency == base_currency:
        return {"amount": float(amount), "currency": base_currency}
    conv = convert(amount, input_currency, base_currency)
    return {
        "amount": conv["converted_amount"],
        "currency": base_currency,
        "original_amount": conv["original_amount"],
        "original_currency": conv["original_currency"],
        "exchange_rate": conv["exchange_rate"],
    }''',
        "amount=500.0, input_currency='INR', base_currency='USD'",
        "{'amount': 5.20, 'currency': 'USD', 'original_amount': 500.0, "
        "'original_currency': 'INR', 'exchange_rate': 0.0104}",
        "frankfurter.dev returns a rate (cached for 1 hour). The original amount + "
        "currency + rate are preserved alongside the converted amount so the UI can "
        "show 'converted from INR 500 @ rate 0.0104'.",
    )

    flow_step(doc,
        "Step 1.7 — Build MongoDB document and insert",
        "routes/ai.py",
        '''doc = {
    "title": parsed["title"],
    "category": parsed["category"],
    "date": date_str,
    **amount_fields,
}
result = await expense_collection.insert_one(doc)''',
        "doc = {\n"
        "  'title': 'Zomato', 'category': 'Food',\n"
        "  'date': '18 May 2026',\n"
        "  'amount': 1000.0, 'currency': 'INR'\n"
        "}",
        "result.inserted_id = ObjectId('6a08...')",
        "Motor's insert_one is awaited (async). MongoDB assigns a 12-byte ObjectId and "
        "returns it. The original `doc` is mutated in place to include the new _id field.",
    )

    flow_step(doc,
        "Step 1.8 — Build and return JSON response",
        "routes/ai.py",
        '''return {
    "message": "Expense parsed and added",
    "expense": {
        "id": str(result.inserted_id),
        "title": doc["title"],
        "category": doc["category"],
        "date": doc["date"],
        **amount_fields,
    },
}''',
        "(internal doc + ObjectId)",
        "{\n"
        '  "message": "Expense parsed and added",\n'
        '  "expense": {\n'
        '    "id": "6a08...",\n'
        '    "title": "Zomato",\n'
        '    "category": "Food",\n'
        '    "date": "18 May 2026",\n'
        '    "amount": 1000.0,\n'
        '    "currency": "INR"\n'
        "  }\n"
        "}",
        "We avoid spreading the full Mongo doc into the response because it contains an "
        "ObjectId which is not JSON-serialisable. Instead we explicitly cast id to str and "
        "spread amount_fields (which contains only primitives).",
    )

    flow_step(doc,
        "Step 1.9 — Streamlit renders success message and reruns (locale-formatted)",
        "frontend/app.py",
        '''from formatting import format_amount

def fmt(amount, currency=None, decimals=0):
    return format_amount(amount, currency or BASE_CURRENCY, decimals)

exp = result["expense"]
msg = (f"Added **{exp['title']}** • {BASE_SYMBOL}{fmt(exp['amount'], decimals=2)} • "
       f"**{exp['category']}** • {exp['date']}")
if "original_currency" in exp and exp["original_currency"] != BASE_CURRENCY:
    msg += (f" _(from {exp['original_currency']} "
            f"{fmt(exp['original_amount'], exp['original_currency'], 2)} "
            f"@ rate {exp['exchange_rate']})_")
st.success(msg)
st.rerun()''',
        "(JSON response from backend)",
        "(Rendered green success bubble in the browser, locale-formatted)\n"
        "Added Shopping mall • $5.20 • Shopping • 22 May 2026\n"
        "(from INR 500.00 @ rate 0.0104)",
        "Amounts are now passed through format_amount() which uses Indian lakh-style "
        "commas for INR (1,00,000) and Western style for everything else (100,000). "
        "st.rerun() then re-fetches /expenses so the new row appears in the list below.",
    )

    # ───────────────────── Flow 2: Chatbot with Tool Use ─────────────
    add_h2(doc, "Flow 2: POST /chat — Chatbot with Tool Use")
    add_para(doc,
        "User scenario: user types 'How much have I spent on food?' on the Chat page. "
        "Groq decides to call get_expenses_by_category('Food'), the backend runs it against "
        "MongoDB and the result is fed back to Groq which writes the natural-language reply."
    )

    flow_step(doc,
        "Step 2.1 — Chat page captures input and posts",
        "frontend/pages/4_Chat.py",
        '''user_input = st.chat_input("Ask me anything about your expenses...")
if user_input:
    st.session_state.chat_history.append({"role": "user", "content": user_input})
    response = post("/chat", {
        "message": user_input,
        "history": st.session_state.chat_history[:-1],
    })
    reply = response.get("reply", ...)
    st.session_state.chat_history.append({"role": "assistant", "content": reply})''',
        "user_input = 'How much have I spent on food?'\n"
        "st.session_state.chat_history = []  (first turn)",
        '{"message": "How much have I spent on food?", "history": []}',
        "Streamlit session_state preserves the chat history across reruns. We send the "
        "history (minus the just-appended user message) so the backend knows the full "
        "conversation context.",
    )

    flow_step(doc,
        "Step 2.2 — FastAPI hands off to the tool-use loop",
        "routes/ai.py",
        '''@router.post("/chat")
def chat_endpoint(req: ChatRequest):
    try:
        history_dicts = [m.dict() for m in (req.history or [])]
        reply = chat_with_tools(req.message, history_dicts)
        return {"reply": reply}
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Chat error: {str(e)}")''',
        "req = ChatRequest(message='...', history=[])",
        "(Calls chat_with_tools(...) → next step)",
        "Notice this is `def`, not `async def`. The Groq SDK is synchronous, so we let "
        "FastAPI run this in its threadpool. Each tool call inside also runs sync (PyMongo).",
    )

    flow_step(doc,
        "Step 2.3 — Tool-use loop, iteration 1: ask Groq",
        "ai/llm_client.py",
        '''messages = [{"role": "system", "content": CHAT_SYSTEM_PROMPT}]
messages.append({"role": "user", "content": user_message})

response = client.chat.completions.create(
    model="openai/gpt-oss-120b",
    messages=messages,
    tools=TOOL_SCHEMAS,
    tool_choice="auto",
    max_tokens=1024,
)
msg = response.choices[0].message
if not msg.tool_calls:
    return msg.content.strip()    # Final text
# else: we have a tool call to execute...''',
        "messages = [\n"
        "  {'role':'system','content':'You are an AI assistant...'},\n"
        "  {'role':'user','content':'How much have I spent on food?'}\n"
        "]\n"
        "tools = [8 tool schemas]",
        "msg.tool_calls = [\n"
        "  {id:'call_x', function:{name:'get_expenses_by_category',\n"
        "                          arguments:'{\"category\":\"Food\"}'}}\n"
        "]",
        "Groq reads the user's question + the descriptions of all 8 tools and decides "
        "get_expenses_by_category is the right one. It returns the function name and a "
        "JSON-encoded argument string.",
    )

    flow_step(doc,
        "Step 2.4 — Execute the tool against MongoDB",
        "ai/tools.py",
        '''def get_expenses_by_category(category: str) -> dict:
    expenses = _find_month_expenses(_current_month_name())
    filtered = [e for e in expenses
                if (e.get("category") or "").lower() == category.lower()]
    total = sum(e.get("amount", 0) for e in filtered)
    return {
        "category": category,
        "count": len(filtered),
        "total": total,
        "items": [{"title": e["title"], "amount": e["amount"]} for e in filtered],
    }''',
        "category = 'Food'",
        "{\n"
        "  'category': 'Food', 'count': 2, 'total': 1450.0,\n"
        "  'items': [\n"
        "    {'title': 'Zomato', 'amount': 1000.0},\n"
        "    {'title': 'Coffee', 'amount': 450.0}\n"
        "  ]\n"
        "}",
        "Sync PyMongo is used here rather than Motor because Groq's tool-call execution "
        "is synchronous. The result is built as a plain dict so it can be JSON-encoded.",
    )

    flow_step(doc,
        "Step 2.5 — Send tool result back to Groq",
        "ai/llm_client.py",
        '''messages.append({"role":"assistant", "content":"", "tool_calls":[...]})
messages.append({
    "role": "tool",
    "tool_call_id": tc.id,
    "name": fn_name,
    "content": json.dumps(result, default=str),
})
# Loop back to client.chat.completions.create(...)''',
        "result dict from Step 2.4",
        "messages = [...previous..., assistant call, tool result]",
        "We append two messages: (a) the assistant's tool-call request and (b) the tool's "
        "result. The next iteration of the for-loop calls Groq again with this enriched "
        "history.",
    )

    flow_step(doc,
        "Step 2.6 — Iteration 2: Groq writes the final answer",
        "ai/llm_client.py",
        '''response = client.chat.completions.create(model=..., messages=messages, tools=...)
msg = response.choices[0].message
if not msg.tool_calls:
    return msg.content.strip()''',
        "messages with tool result included",
        "'You\\'ve spent Rs.1,450 on food this month across 2 entries: "
        "Zomato (Rs.1,000) and Coffee (Rs.450).'",
        "Now that Groq has the data, it stops calling tools and returns a plain text "
        "response (stop_reason='stop' or similar). We return this string.",
    )

    flow_step(doc,
        "Step 2.7 — Display reply in chat bubble",
        "frontend/pages/4_Chat.py",
        '''with st.chat_message("assistant"):
    st.markdown(reply)
st.session_state.chat_history.append({"role": "assistant", "content": reply})''',
        "reply: 'You\\'ve spent Rs.1,450 on food this month...'",
        "(Rendered chat bubble with the assistant's reply)",
        "The reply is appended to chat_history so it persists across reruns. Streamlit's "
        "st.chat_message context manager handles the visual bubble styling.",
    )

    # ───────────────────── Flow 3: Currency switch / bulk conversion ──
    add_h2(doc, "Flow 3: PUT /settings — Switch base currency and bulk-convert all data")

    add_para(doc,
        "User scenario: the user is currently using INR as the base currency and has "
        "many expenses + a salary stored. They switch to USD via the Settings page. The "
        "backend converts every stored amount to USD using the current INR→USD rate."
    )

    flow_step(doc,
        "Step 3.1 — Settings page sends PUT",
        "frontend/pages/5_Settings.py",
        '''new_code = codes[options.index(new_choice)]
if new_code != current_base:
    if st.button(f"Convert everything to {new_code}", type="primary"):
        result = put("/settings", {"base_currency": new_code})
        st.success(result["message"])
        st.info(f"Rate used: 1 {result['old_base_currency']} = "
                f"{result['conversion_rate']} {result['base_currency']}")
        st.rerun()''',
        "new_code = 'USD' (current_base = 'INR')",
        '{"base_currency": "USD"}',
        "The Streamlit page builds a dropdown of all supported currencies. When the chosen "
        "code differs from the current base, a primary button appears.",
    )

    flow_step(doc,
        "Step 3.2 — Backend fetches old base + exchange rate",
        "routes/settings.py",
        '''doc = await settings_collection.find_one({"_id": SETTINGS_ID})
old_base = (doc or {}).get("base_currency", DEFAULT_BASE_CURRENCY)
rate = get_rate(old_base, new_base)''',
        "settings.base_currency='USD' (from request)",
        "old_base = 'INR', rate = 0.01044  (1 INR = 0.01044 USD)",
        "We read the previous base from Mongo. get_rate calls Frankfurter (cached for 1 "
        "hour) to find the conversion factor from old→new.",
    )

    flow_step(doc,
        "Step 3.3 — Bulk-update all expenses",
        "routes/settings.py",
        '''expenses = await expense_collection.find().to_list(10000)
for exp in expenses:
    new_amount = round(exp.get("amount", 0) * rate, 2)
    await expense_collection.update_one(
        {"_id": exp["_id"]},
        {"$set": {"amount": new_amount, "currency": new_base}},
    )''',
        "expenses = [ {_id:..., amount: 1000, currency:'INR'}, ... ]",
        "(MongoDB documents updated in place)\n"
        "Each expense: amount: 1000 → 10.44, currency: 'INR' → 'USD'",
        "Each expense is updated individually with $set. For a personal-use dataset (a "
        "few hundred rows) this is fine; for thousands of rows you would use bulk_write.",
    )

    flow_step(doc,
        "Step 3.4 — Same for salaries, then persist new base",
        "routes/settings.py",
        '''await settings_collection.update_one(
    {"_id": SETTINGS_ID},
    {"$set": {"base_currency": new_base}},
    upsert=True,
)
return {
    "message": f"Base currency changed from {old_base} to {new_base}...",
    "base_currency": new_base,
    "old_base_currency": old_base,
    "conversion_rate": rate,
    "expenses_converted": len(expenses),
    "salaries_converted": len(salaries),
}''',
        "(end of bulk loops)",
        '{"message": "...", "base_currency": "USD", "old_base_currency": "INR", '
        '"conversion_rate": 0.01044, "expenses_converted": 47, "salaries_converted": 1}',
        "The settings document is upserted (insert-or-update) so the change survives "
        "process restart. The response includes counts and the rate used for "
        "transparency.",
    )

    # ───────────────────── Flow 4: Insights ───────────────────────────
    add_h2(doc, "Flow 4: GET /insights/monthly — AI-written spending paragraph")

    flow_step(doc,
        "Step 4.1 — User clicks button on Insights page",
        "frontend/pages/3_AI_Insights.py",
        '''if st.button("Generate Insights", key="insights_btn"):
    with st.spinner("AI is analyzing your spending..."):
        data = get("/insights/monthly")
        st.write(data.get("insights", "No insights available."))''',
        "(Click event)",
        "(HTTP request)\nGET /insights/monthly",
        "Streamlit shows a spinner during the network round-trip. No request body for GET.",
    )

    flow_step(doc,
        "Step 4.2 — Backend assembles the monthly summary",
        "routes/ai.py",
        '''async def _build_monthly_summary() -> dict:
    current_month_name = datetime.now().strftime("%B")
    current_month_year = datetime.now().strftime("%B-%Y")
    expenses = await fetch_expenses_by_month(current_month_name)
    summary = {}
    total = 0
    for e in expenses:
        cat = e.get("category", "Miscellaneous")
        summary[cat] = summary.get(cat, 0) + e.get("amount", 0)
        total += e.get("amount", 0)
    return {"month": current_month_year, "summary": summary,
            "total": total, "currency": "INR"}''',
        "(no inputs — uses datetime.now())",
        "{'month': 'May-2026', 'summary': {'Food': 1450, 'Shopping': 3000}, 'total': 4450}",
        "Reuses the regex-based fetch_expenses_by_month helper from routes/expenses.py. "
        "Groups expenses by category and computes a total.",
    )

    flow_step(doc,
        "Step 4.3 — Send summary to Groq with INSIGHTS_PROMPT",
        "ai/llm_client.py",
        '''def generate_monthly_insights(summary_data: dict, salary):
    user_msg = json.dumps({
        "month": summary_data.get("month"),
        "total_spent": summary_data.get("total"),
        "category_breakdown": summary_data.get("summary"),
        "salary": salary,
        "currency": summary_data.get("currency", "INR"),
    }, indent=2)
    return _ask_groq(INSIGHTS_PROMPT, user_msg, max_tokens=400)''',
        "summary_data (from step 4.2), salary = 70000",
        "'In May 2026 you spent Rs.4,450 out of your Rs.70,000 salary, which means you "
        "are saving 94 percent of your income — excellent. Your biggest category was "
        "Shopping at 67 percent...'",
        "Groq is called with INSIGHTS_PROMPT as the system prompt. The user message is "
        "the spending data as JSON.",
    )

    # ───────────────────── Flow 5: Monthly summary with filter ────────
    add_h2(doc, "Flow 5: GET /expenses/summary/monthly?month=April&year=2026")

    flow_step(doc,
        "Step 5.1 — Monthly Summary page builds the URL",
        "frontend/pages/1_Monthly_Summary.py",
        '''selected_month = st.selectbox("Month", MONTHS, ...)
selected_year = st.selectbox("Year", year_options, ...)
data = get(f"/expenses/summary/monthly?month={selected_month}&year={selected_year}")''',
        "selected_month='April', selected_year=2026",
        "GET /expenses/summary/monthly?month=April&year=2026",
        "Query parameters are URL-encoded directly into the path.",
    )

    flow_step(doc,
        "Step 5.2 — Backend filters by month + year",
        "routes/expenses.py",
        '''@router.get("/expenses/summary/monthly")
async def get_monthly_summary(month: str = None, year: int = None):
    if month and year:
        target_month_name = month
        target_month_year = f"{month}-{year}"
    else:
        target_month_name = datetime.now().strftime("%B")
        target_month_year = datetime.now().strftime("%B-%Y")
    expenses = await fetch_expenses_by_month(target_month_name)
    if year:
        expenses = [e for e in expenses if str(year) in str(e.get("date", ""))]''',
        "month='April', year=2026",
        "expenses = [list of expense dicts from April 2026 only]",
        "FastAPI auto-parses query params as function arguments using type hints. We then "
        "filter the regex-matched results to only those whose date string contains the "
        "year. This is a simplification — the date is stored as '15 April 2026' so the "
        "year filter is a substring match.",
    )

    flow_step(doc,
        "Step 5.3 — Streamlit renders pie chart + table (locale-aware)",
        "frontend/pages/1_Monthly_Summary.py",
        '''def fmt(amount, decimals=0):
    return format_amount(amount, BASE_CURRENCY, decimals)

df = pd.DataFrame(
    [{"Category": k, "Amount": v} for k, v in data["summary"].items()]
).sort_values("Amount", ascending=False)
df["Amount_fmt"] = df["Amount"].apply(lambda v: fmt(v))

fig = px.pie(df, names="Category", values="Amount", hole=0.4)
fig.update_traces(
    customdata=df[["Amount_fmt"]].values,
    hovertemplate=f"%{{label}}: {SYMBOL}%{{customdata[0]}}<extra></extra>",
)
st.plotly_chart(fig, use_container_width=True)''',
        '{"month":"April-2026","summary":{"Food":3000.0,"Shopping":500.0}, ...}',
        "(Interactive pie chart with Indian-style commas on hover)",
        "Plotly's built-in number formatting doesn't support Indian locale, so we "
        "pre-compute formatted strings in a column and pass them via custom_data. "
        "The hovertemplate references customdata[0] instead of the raw value.",
    )

    # ───────────────────── Flow 6: Chatbot — "last month" ─────────────
    add_h2(doc, "Flow 6: POST /chat — \"How much did I spend last month?\"")
    add_para(doc,
        "User scenario: today is 22 May 2026; user types 'How much did I spend last "
        "month?' on the Chat page. The LLM has to (a) resolve 'last month' to "
        "April 2026 using today's date injected into the prompt, then (b) call the "
        "month-aware tool, then (c) reply using the pre-formatted display string."
    )

    flow_step(doc,
        "Step 6.1 — Backend injects today's date into the system prompt",
        "ai/llm_client.py",
        '''def chat_with_tools(user_message: str, history: list) -> str:
    now = datetime.now()
    system_prompt = CHAT_SYSTEM_PROMPT.format(
        today=now.strftime("%d %B %Y"),
        current_month=now.strftime("%B"),
        current_year=now.year,
    )
    messages = [{"role": "system", "content": system_prompt}]
    ...''',
        "user_message = 'How much did I spend last month?'\n"
        "today = 22 May 2026",
        "system_prompt includes:\n"
        "  'Today's date is: 22 May 2026'\n"
        "  'Current month: May 2026'",
        "Without injecting today's date the LLM doesn't reliably know what 'last month' "
        "means — it might use its training-cutoff date. With the date in the system "
        "prompt the model can do simple arithmetic: May - 1 = April.",
    )

    flow_step(doc,
        "Step 6.2 — LLM picks the right month-aware tool",
        "ai/tools.py + TOOL_SCHEMAS",
        '''TOOL_SCHEMAS = [
    ...,
    {"type": "function", "function": {
        "name": "get_expenses_for_month",
        "description": "Get all expenses for a specific month and year...",
        "parameters": {
            "type": "object",
            "properties": {
                "month": {"type": "string", "description": "Full month name..."},
                "year": {"type": "integer", "description": "4-digit year..."},
            },
            "required": ["month", "year"],
        },
    }},
    ...
]''',
        "(LLM reasoning, not directly visible)",
        "tool_call: get_expenses_for_month(month='April', year=2026)",
        "Three month-aware tools are exposed: get_yearly_summary, "
        "get_expenses_for_month, get_expenses_by_category_for_month. The LLM picks "
        "based on the question shape — a 'how much last month' question selects "
        "get_expenses_for_month with the month string + year integer.",
    )

    flow_step(doc,
        "Step 6.3 — Tool queries MongoDB and pre-formats the result",
        "ai/tools.py",
        '''def get_expenses_for_month(month: str, year: int) -> dict:
    cur = _base_currency()
    expenses = _find_month_expenses(month)
    expenses = [e for e in expenses if str(year) in str(e.get("date", ""))]
    total = sum(e.get("amount", 0) for e in expenses)
    return {
        "month": month, "year": year, "currency": cur,
        "count": len(expenses),
        "total": total,
        "total_display": _fmt(total, 2, cur),     # <-- pre-formatted
        "expenses": [
            {"title": e.get("title"), "amount": e.get("amount"),
             "amount_display": _fmt(e.get("amount"), 2, cur),
             "category": e.get("category"), "date": e.get("date")}
            for e in expenses
        ],
    }''',
        "month='April', year=2026",
        "{\n"
        "  'month': 'April', 'year': 2026, 'currency': 'INR',\n"
        "  'count': 5, 'total': 355033.17,\n"
        "  'total_display': 'Rs.3,55,033.17',\n"
        "  'expenses': [{'title': 'Zomato', 'amount': 1000.0,\n"
        "                'amount_display': 'Rs.1,000.00', ...}, ...]\n"
        "}",
        "The tool reads the base currency from MongoDB settings, applies the same "
        "regex-based month filter as the HTTP endpoint, and pre-formats every amount "
        "using services.currency.format_amount (Indian lakh style for INR). The raw "
        "numeric 'total' is kept for downstream calculation; 'total_display' is the "
        "string the LLM will copy into its reply.",
    )

    flow_step(doc,
        "Step 6.4 — LLM writes the reply using the _display field",
        "ai/prompts.py (CHAT_SYSTEM_PROMPT)",
        '''CRITICAL — formatting amounts in your reply:
- Every tool result that contains amounts ALSO contains pre-formatted
  versions in keys ending in '_display'...
- When mentioning an amount in your reply, USE the _display string
  EXACTLY as given. Do NOT reformat numbers, do NOT add or change the
  currency symbol, do NOT round the number yourself.''',
        "(tool result from Step 6.3)",
        "reply: 'You spent Rs.3,55,033.17 in April 2026.'",
        "The system prompt forces the LLM to copy total_display verbatim. This is a "
        "common AI-engineering trick: rather than asking the model to format numbers "
        "(unreliable), pre-format on the server and have the model copy the string. "
        "Same approach as JSON-mode for structured output.",
    )

    doc.add_page_break()


# ────────────────────────────────────────────────────────────────────────────
# Section 5: Architecture diagram
# ────────────────────────────────────────────────────────────────────────────


def section_5_architecture(doc):
    add_h1(doc, "5. Architecture Diagram")
    add_para(doc,
        "The diagram below shows every layer of the system, the files that live in "
        "each layer and the external services involved. Each arrow is labelled with "
        "what type of data flows across it."
    )
    diagram = r"""
+--------------------------------------------------------------------------+
|                            USER (Web Browser)                            |
|                       Chrome / Safari / Firefox                          |
+----------------------------------+---------------------------------------+
                                   |   HTML, CSS, JS, WebSocket
                                   |   user clicks, form submits
                                   v
+--------------------------------------------------------------------------+
|                      STREAMLIT SERVER (port 8501)                        |
|                                                                          |
|  +------------------+      +-------------------------------+             |
|  | frontend/app.py  | <--> | frontend/pages/1_*.py to 5_*  |             |
|  |  (Dashboard)     |      |  (Monthly, Yearly, Insights,  |             |
|  +------------------+      |   Chat, Settings)             |             |
|         |                  +-------------------------------+             |
|         |  uses                              |  all import               |
|         v                                    v                           |
|  +-------------------------------------------------------------+         |
|  |                  frontend/api_client.py                     |         |
|  |          get / post / put / delete -> httpx                 |         |
|  +-------------------------------------------------------------+         |
+----------------------------------+---------------------------------------+
                                   |  JSON over HTTP
                                   |  e.g. POST /expenses/natural
                                   v
+--------------------------------------------------------------------------+
|                       FASTAPI BACKEND (port 8000)                        |
|                                                                          |
|  +---------------+   +---------------+   +-----------+   +-----------+   |
|  | routes/       |   | routes/       |   | routes/   |   | routes/   |   |
|  |  expenses.py  |   |  salary.py    |   |  ai.py    |   | settings  |   |
|  +-------+-------+   +-------+-------+   +-----+-----+   +-----+-----+   |
|          |                   |                 |               |         |
|          +-----+-------------+-----------------+---------------+         |
|                |                                                         |
|                |  uses (Pydantic validation, ObjectId, datetime)         |
|                v                                                         |
|  +-------------------------------------------------------------+         |
|  |  models.py        <----   used by every route               |         |
|  +-------------------------------------------------------------+         |
|                |                              |                          |
|  uses for DB   |                              |  uses for AI             |
|                v                              v                          |
|  +-----------------+      +----------------------------------+           |
|  |  database.py    |      |  ai/llm_client.py                |           |
|  |  (Motor async)  |      |     +----------------------+     |           |
|  |  collections:   |      |     | uses ai/prompts.py   |     |           |
|  |   expenses,     |      |     | uses ai/tools.py     |     |           |
|  |   salaries,     |      |     +----------------------+     |           |
|  |   settings      |      +-----------------+----------------+           |
|  +--------+--------+                        |                            |
|           |                                 |  uses for FX               |
|           |                                 v                            |
|           |              +------------------------------+                |
|           |              |  services/currency.py        |                |
|           |              |  (httpx, in-memory cache)    |                |
|           |              +-------------+----------------+                |
|           |                            |                                 |
+-----------|----------------------------|---------------------------------+
            |                            |
            |  TLS over the internet     |  HTTPS GET
            v                            v
+---------------------+      +-------------------------+    +-------------+
|  MongoDB Atlas      |      |  Groq API               |    | Frankfurter |
|  (cloud database)   |      |  openai/gpt-oss-120b    |    | (ECB FX     |
|  bson docs          |      |  Tool Use / Function    |    |  rates)     |
|  expenses,          |      |  Calling protocol       |    | JSON        |
|  salaries,          |      |                         |    +-------------+
|  settings           |      +-------------------------+
+---------------------+
"""
    add_code(doc, diagram)
    doc.add_page_break()


def section_6_design_decisions(doc):
    add_h1(doc, "6. Key Design Decisions")

    # ----
    add_h2(doc, "Decision: Split frontend and backend into two processes")
    add_para(doc,
        "Streamlit alone CAN talk to MongoDB directly, but combining them would put DB "
        "credentials and the Groq API key inside the frontend process — anything in the "
        "Streamlit script becomes visible to anyone who can hit the page. Splitting also "
        "lets the FastAPI backend serve a future mobile app or WhatsApp bot from the same "
        "API surface."
    )

    # ----
    add_h2(doc, "Decision: Provider-agnostic LLM layer in ai/llm_client.py")
    add_para(doc,
        "Every Groq call goes through ai/llm_client.py. Inside that file we use Groq's "
        "OpenAI-compatible API (chat.completions.create). Swapping providers — to Claude, "
        "OpenAI proper, Gemini — is a one-file change. This is why we expose four "
        "high-level functions (generate_monthly_insights, generate_budget_advice, "
        "categorize_expense, parse_expense_text, chat_with_tools) rather than letting "
        "routes call the SDK directly."
    )

    # ----
    add_h2(doc, "Decision: Async (Motor) for HTTP routes, sync (PyMongo) inside AI tools")
    add_para(doc,
        "FastAPI is async-first and Motor lets us await Mongo from inside endpoints, so "
        "thousands of concurrent requests don't block on I/O. However, the Groq SDK is "
        "synchronous and the tool-use loop (chat_with_tools) calls our 8 tool functions "
        "inline. Forcing those tools to be async would mean wrapping every DB call in "
        "asyncio.run() (an anti-pattern) or using nest_asyncio (a hack). Instead we open "
        "a second, sync PyMongo client inside ai/tools.py. The two clients share the same "
        "MongoDB Atlas cluster but operate independently."
    )
    add_para(doc, "Affected modules:")
    add_bullet(doc, "database.py  — Motor (async)")
    add_bullet(doc, "ai/tools.py  — PyMongo (sync)")

    # ----
    add_h2(doc, "Decision: Currency stored in base, original kept for transparency")
    add_para(doc,
        "When the user adds a USD expense with INR as the base currency, we convert to "
        "INR at insert time. Crucially we ALSO store original_amount, original_currency "
        "and exchange_rate fields in the same document. This means future totals are "
        "always in base currency (no live conversion at read time), but we never lose "
        "the original transaction value. When the base currency changes, all amounts are "
        "bulk-converted but the original_* fields are preserved as historical reference."
    )

    # ----
    add_h2(doc, "Decision: In-memory exchange-rate cache with 1-hour TTL")
    add_para(doc,
        "frankfurter.dev is free and key-less but they politely ask you not to hammer it. "
        "services/currency.py keeps a dict of (from->to: rate) and only re-fetches if more "
        "than 3600 seconds have passed. This is intentionally simple (one process, no "
        "shared cache). For multi-instance deployment we would move this to Redis."
    )

    # ----
    add_h2(doc, "Decision: Pydantic models defined once, used three ways")
    add_para(doc,
        "Every class in models.py simultaneously serves as request validation, OpenAPI "
        "schema (visible at /docs) and the type used inside route handlers. Pydantic v2 "
        "enforces types at the boundary; inside the handler the parameter is a "
        "fully-typed object."
    )

    # ----
    add_h2(doc, "Decision: 8 tool functions for the chatbot")
    add_para(doc,
        "Instead of giving the LLM a single mega-tool (e.g. run_mongo_query), we expose "
        "eight tightly-scoped functions: get_all_expenses_this_month, "
        "get_expenses_by_category, get_monthly_total, get_salary, get_savings, "
        "add_expense, delete_expense_by_title, update_expense_by_title. Each has a "
        "well-typed signature plus a docstring that the LLM uses to choose between them. "
        "This dramatically improves tool-call reliability — the LLM doesn't have to "
        "compose queries; it just picks the right pre-built function."
    )

    # ----
    add_h2(doc, "Decision: Force JSON-only output for parse_expense_text")
    add_para(doc,
        "PARSE_EXPENSE_PROMPT contains the phrase 'Return ONLY a valid JSON object'. "
        "Despite that, models sometimes wrap output in markdown ```json ... ``` fences. "
        "parse_expense_text strips those fences before json.loads. After parsing it "
        "validates that the returned category is in the allowed set; otherwise it falls "
        "back to 'Miscellaneous'. Two-layer defence."
    )

    # ----
    add_h2(doc, "Decision: Reuse fetch_expenses_by_month across modules")
    add_para(doc,
        "fetch_expenses_by_month is defined once in routes/expenses.py and imported by "
        "routes/salary.py (for savings), routes/ai.py (for summaries and insights) and "
        "ai/tools.py uses the same regex pattern. This is intentional DRY: the date "
        "format ('15 May 2026') is application-specific and the matching logic lives in "
        "exactly one place."
    )

    # ----
    add_h2(doc, "Decision: Streamlit multi-page layout via filename prefixes")
    add_para(doc,
        "Streamlit auto-detects files inside frontend/pages/ and orders them by filename. "
        "We use a numeric prefix (1_, 2_, 3_, ...) so the sidebar order is deterministic. "
        "The home page is frontend/app.py — Streamlit treats whatever you `streamlit run` "
        "as the entry."
    )

    # ----
    add_h2(doc, "Decision: AI-detected currency in natural-language input, with 3-tier fallback")
    add_para(doc,
        "Earlier, the Natural Language tab required a separate 'Currency typed in' dropdown — "
        "the user had to manually pick the currency that matched what they typed. This was "
        "redundant since the model can see the same text. PARSE_EXPENSE_PROMPT was extended "
        "to ask for a 4th key (currency), with mappings for natural-language hints "
        "(rs/rupees -> INR, $/dollars -> USD, etc.). parse_expense_text validates the returned "
        "code against the supported set. The endpoint then uses a 3-tier fallback chain: "
        "AI-detected currency, explicit currency parameter, base currency. The dropdown was "
        "removed from the UI; users now just type."
    )

    # ----
    add_h2(doc, "Decision: Pre-format display strings in tool results, not in LLM output")
    add_para(doc,
        "When the chatbot answers 'how much in April', it needs to write a number with the "
        "right locale formatting (Rs.3,55,033 for INR, $355,033 for USD). Telling the LLM "
        "'format Indian-style for INR' in the system prompt works ~70 percent of the time — "
        "the model occasionally forgets or uses Western style. Instead we pre-format every "
        "amount in the tool result as *_display strings (total_display, amount_display, "
        "salary_display, monthly_totals_display, ...) and the system prompt strictly tells "
        "the LLM to copy the _display value verbatim. Reliability jumps to ~100 percent. "
        "Same principle as forced JSON output for structured extraction: don't ask the model "
        "to format, ask it to copy."
    )

    # ----
    add_h2(doc, "Decision: Inject today's date into the chat system prompt")
    add_para(doc,
        "Users say things like 'last month' or 'this year' which the model needs to resolve "
        "to concrete month + year arguments. Models have no built-in clock — they may use "
        "training-cutoff dates. chat_with_tools() formats the system prompt at request time "
        "with today's full date, current month name, and current year. The model can then "
        "reliably compute 'last month' = current month - 1."
    )

    # ----
    add_h2(doc, "Decision: Month-aware tools alongside current-month tools")
    add_para(doc,
        "The original chatbot only had current-month tools (get_all_expenses_this_month, "
        "get_monthly_total, get_expenses_by_category) so questions about past months failed. "
        "We could have given the existing tools optional month/year parameters, but the LLM "
        "tends to handle clearly-named single-purpose tools better than overloaded ones. "
        "Instead, three new tools were added: get_yearly_summary, get_expenses_for_month "
        "(month, year), get_expenses_by_category_for_month (category, month, year). The "
        "system prompt explicitly tells the LLM which tool to pick for which question shape."
    )

    # ----
    add_h2(doc, "Decision: Validation errors map to HTTP 400, not 500")
    add_para(doc,
        "When testing the API via Swagger UI's 'Try it out', the default body has "
        "'currency': 'string' as a placeholder. Hitting submit without changing it used to "
        "trigger Frankfurter to error on the bogus code, then propagate as a 500 from our "
        "server. build_amount_fields and _to_base_currency now raise ValueError early for "
        "any unsupported code; every expense + salary route catches that and returns HTTP "
        "400 with a clear message listing the supported codes. Same change pattern was "
        "applied to the AI smart-add and natural-add endpoints."
    )

    # ----
    add_h2(doc, "Decision: Locale-aware number formatting via a shared helper")
    add_para(doc,
        "format_amount in services/currency.py implements both the Indian lakh comma style "
        "(1,00,000) and the Western style (100,000). The frontend has a thin copy at "
        "frontend/formatting.py (same logic) because Streamlit's working directory makes "
        "importing from services/ awkward. The same helper is imported by ai/tools.py for "
        "the chatbot's _display strings, by frontend/app.py for the Dashboard, and by the "
        "Monthly/Yearly Summary pages for chart hover and table cells. Plotly's built-in "
        "number formatting doesn't support Indian locale, so the bar/pie chart hover labels "
        "are precomputed strings passed via custom_data."
    )


# ────────────────────────────────────────────────────────────────────────────
def main():
    doc = Document()
    configure_document(doc)
    cover_page(doc)
    section_1_project_overview(doc)
    section_2_entry_points(doc)
    section_3_dependency_map(doc)
    section_4_data_flow(doc)
    section_5_architecture(doc)
    section_6_design_decisions(doc)

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(f"Wrote: {OUTPUT}")


if __name__ == "__main__":
    main()
