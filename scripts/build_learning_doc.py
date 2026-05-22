"""Build a Word document containing the architecture guide + every source file.

Run from project root:
    source venv/bin/activate
    python scripts/build_learning_doc.py

Output: docs/AI_Expense_Tracker_Learning_Guide.docx
"""
import os
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

PROJECT_ROOT = Path(__file__).parent.parent
ARCHITECTURE_MD = PROJECT_ROOT / "docs" / "ARCHITECTURE.md"
OUTPUT_DOCX = PROJECT_ROOT / "docs" / "AI_Expense_Tracker_Learning_Guide.docx"


# Files to include in the "Appendix: All Source Files" section
SOURCE_FILES = [
    ("Backend Entry & Database", [
        "main.py",
        "database.py",
        "models.py",
        "requirements.txt",
    ]),
    ("Routes", [
        "routes/__init__.py",
        "routes/expenses.py",
        "routes/salary.py",
        "routes/ai.py",
        "routes/settings.py",
    ]),
    ("AI Package", [
        "ai/__init__.py",
        "ai/prompts.py",
        "ai/tools.py",
        "ai/llm_client.py",
    ]),
    ("Services", [
        "services/__init__.py",
        "services/currency.py",
    ]),
    ("Frontend", [
        "frontend/app.py",
        "frontend/api_client.py",
        "frontend/formatting.py",
        "frontend/pages/1_Monthly_Summary.py",
        "frontend/pages/2_Yearly_Summary.py",
        "frontend/pages/3_AI_Insights.py",
        "frontend/pages/4_Chat.py",
        "frontend/pages/5_Settings.py",
    ]),
    ("Project Files", [
        "README.md",
        ".gitignore",
    ]),
]


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    return h


def add_paragraph(doc, text, bold=False, italic=False, size=11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    return p


def add_code_block(doc, code: str, language: str = ""):
    """Add a monospace code block with light grey background-ish styling."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.2)
    run = p.add_run(code)
    run.font.name = "Courier New"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    return p


def parse_markdown_and_add(doc, md_text: str):
    """Parse a subset of markdown and add to docx.

    Supports: headings (#, ##, ###), code fences (```), bullet lists, plain paragraphs.
    """
    lines = md_text.split("\n")
    i = 0
    in_code = False
    code_buffer = []
    code_lang = ""

    while i < len(lines):
        line = lines[i]

        # Code fence
        if line.startswith("```"):
            if in_code:
                # Close fence
                add_code_block(doc, "\n".join(code_buffer), code_lang)
                code_buffer = []
                in_code = False
            else:
                # Open fence
                code_lang = line[3:].strip()
                in_code = True
            i += 1
            continue

        if in_code:
            code_buffer.append(line)
            i += 1
            continue

        # Headings
        if line.startswith("# "):
            add_heading(doc, line[2:].strip(), level=1)
        elif line.startswith("## "):
            add_heading(doc, line[3:].strip(), level=2)
        elif line.startswith("### "):
            add_heading(doc, line[4:].strip(), level=3)
        elif line.startswith("#### "):
            add_heading(doc, line[5:].strip(), level=4)

        # Horizontal rule
        elif line.strip() == "---":
            doc.add_paragraph("─" * 60)

        # Bullet list
        elif line.startswith("- ") or line.startswith("* "):
            doc.add_paragraph(line[2:].strip(), style="List Bullet")

        # Numbered list
        elif len(line) > 2 and line[0].isdigit() and line[1:3] in (". ", ") "):
            doc.add_paragraph(line[3:].strip(), style="List Number")

        # Blockquote
        elif line.startswith("> "):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.3)
            run = p.add_run(line[2:].strip())
            run.italic = True

        # Empty line
        elif line.strip() == "":
            pass  # skip

        # Plain paragraph (handles inline **bold** rudimentarily)
        else:
            p = doc.add_paragraph()
            # very rudimentary inline markup handling
            remaining = line
            while "**" in remaining:
                pre, _, after = remaining.partition("**")
                if pre:
                    p.add_run(pre)
                bold_text, _, after2 = after.partition("**")
                if bold_text:
                    r = p.add_run(bold_text)
                    r.bold = True
                remaining = after2
            if remaining:
                p.add_run(remaining)

        i += 1


def add_source_files_appendix(doc):
    add_heading(doc, "APPENDIX: Every Source File", level=1)
    add_paragraph(
        doc,
        "This appendix contains the full source code of every file in the project. "
        "Files are organised by their role in the architecture.",
        italic=True,
    )

    for section_title, file_list in SOURCE_FILES:
        add_heading(doc, section_title, level=2)
        for relpath in file_list:
            file_path = PROJECT_ROOT / relpath
            add_heading(doc, relpath, level=3)
            if not file_path.exists():
                add_paragraph(doc, "(file not found — may be empty or omitted)",
                              italic=True)
                continue
            try:
                content = file_path.read_text()
            except Exception as e:
                add_paragraph(doc, f"(could not read file: {e})", italic=True)
                continue
            if not content.strip():
                add_paragraph(doc, "(empty file — used as Python package marker)",
                              italic=True)
                continue
            # Detect language for syntax highlighting hint
            if relpath.endswith(".py"):
                lang = "python"
            elif relpath.endswith(".md"):
                lang = "markdown"
            elif relpath.endswith(".txt"):
                lang = "text"
            else:
                lang = ""
            add_code_block(doc, content, lang)


def main():
    doc = Document()

    # ---- Title page ----
    title = doc.add_heading("AI Expense Tracker", level=0)
    subtitle = doc.add_paragraph()
    s_run = subtitle.add_run("Complete Architecture & Learning Guide")
    s_run.italic = True
    s_run.font.size = Pt(14)

    doc.add_paragraph()
    author = doc.add_paragraph()
    author_run = author.add_run("Built by Hiya Arora")
    author_run.bold = True

    doc.add_page_break()

    # ---- Architecture content from markdown ----
    if ARCHITECTURE_MD.exists():
        md_text = ARCHITECTURE_MD.read_text()
        parse_markdown_and_add(doc, md_text)
    else:
        add_paragraph(doc, "ARCHITECTURE.md not found", italic=True)

    doc.add_page_break()

    # ---- Appendix: every source file ----
    add_source_files_appendix(doc)

    # ---- Save ----
    OUTPUT_DOCX.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT_DOCX)
    print(f"Wrote: {OUTPUT_DOCX}")


if __name__ == "__main__":
    main()
